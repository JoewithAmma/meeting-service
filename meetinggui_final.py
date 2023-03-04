# -*- coding: utf-8 -*-
"""
Created on Mon Jan 16 08:59:12 2023

@author: User
"""
import PySimpleGUI as sg
import os
import requests
from docx import Document
from requests_toolbelt import MultipartEncoder
#####開始撰寫GUI介面
#畫面顏色主題


sg.theme("LightBlue2")
layout = [[sg.Text("您好，歡迎使用會議記錄小工具")],
          [sg.Text("請輸入會議名稱")],
          [sg.Input()],
          [sg.Text('請選擇影音檔案路徑：',  size=(50, 1))],
          [sg.Input('  ', key="_FILE_", readonly=True,size=(60, 1)),sg.FileBrowse('選擇路徑')],
          [sg.Text('請選擇輸出資料夾路徑：',  size=(50, 1))],
          [sg.Input('  ', key="_FOLDER_", readonly=True,size=(60, 1)),sg.FolderBrowse('選擇路徑')],
          [sg.Button('開始執行'),sg.Exit("離開")]]
#建立窗口
window = sg.Window('會議記錄小工具', layout) 
#讀取窗口事件及訊息
event,values = window.read()
window.close()
#按下離開鍵的事件
if(event=='離開'):
    os._exit(0)
#選取影片或音檔路徑
file_position=values["選擇路徑"]
#輸出的資料夾路徑
output_position=values["選擇路徑0"]
#把檔名切割出來，除了副檔名以外
file_name = file_position.split('/')[-1].split('\\')[-1]

file_dir=file_position.replace(file_name,"")
#程式中止狀況
if(values[0]=="" or values["選擇路徑"]==""):
    sg.popup("會議名稱尚未輸入或檔案未選取，程式終止")
    os._exit(0)
#開始轉換檔案格式    
sg.popup("開始進行語音辨識，系統會於背景執行，執行完成會提示您")
os.chdir(file_dir)   #修改当前工作目录
print(os.getcwd())

#串接API
#https://4957-2404-0-823b-a65a-7da0-4a1-be0-561b.jp.ngrok.io/upload
#http://127.0.0.1:8000/upload       
#url = 'http://127.0.0.1:8000/upload'
url = 'https://be12-2404-0-8237-bcd6-dbd-c98a-75b-f38a.jp.ngrok.io/upload'
filename = file_name
m = MultipartEncoder(fields={'file': (filename, open(filename, 'rb'))})
resp = requests.post(url=url, data=m, headers={'Content-Type': m.content_type})
print(resp.request.headers)
#傳回API連接到AZURE核心的內容，是以陣列方式儲存
text=resp.json()
#開始將逐字稿撰寫到WORD檔裡面
document = Document()
document.add_heading(file_name+'會議記錄逐字稿', level=1)
document.add_paragraph(text[0])
document.save(output_position+"/"+file_name+".docx")
#撰寫文字+時間軸檔案
with open(output_position+"/"+file_name+".srt", "w") as f:
    f.write(text[1])
#撰寫逐字稿文字檔        
with open(output_position+"/"+file_name+".txt","w+") as f:
    f.write(text[0])
    
document = Document()
document.add_heading(file_name+'會議記錄摘要', level=1)
document.add_paragraph(text[2])
document.save(output_position+"/"+file_name+"摘要.docx")
sg.popup("會議紀錄逐字稿及摘要整理完成")
