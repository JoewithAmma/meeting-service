import PySimpleGUI as sg
import azure.cognitiveservices.speech as speechsdk
import time
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import os
import subprocess


#GUI排版
layout = [[sg.Text("您好，歡迎使用會議記錄小工具")],
          [sg.Text("請輸入會議名稱")],
          [sg.Input()],
          [sg.Text('請選擇影音檔案路徑：',  size=(50, 1))],
          [sg.Input('  ', key="_FILE_", readonly=True,size=(60, 1)),sg.FileBrowse('選擇路徑')],
          [sg.Text("請輸入要特別標記的字詞")],
          [sg.Input()],
          [sg.Button('開始執行'),sg.Exit("離開")]]

#建立窗口
window = sg.Window('會議記錄小工具', layout)  # Part 3 - 窗口定义

#讀取窗口事件及訊息
event,values = window.read()
window.close()

#按下離開鍵的事件
if(event=='離開'):
    os._exit(0)
file_position=values[1]["選擇路徑"]

print(values[1][1])


sg.popup("開始進行轉檔")
#獲取檔案名稱
file_name = file_position.split('/')[-1].split('\\')[-1].split('.')[0]
#開始進行任何檔案的轉檔，轉換成wav

subprocess.call(['./ffmpeg/bin/ffmpeg.exe', '-i', file_position,
                   file_name+'.wav'])

sg.popup('檔案格式轉換完成')


#連續的Azure語音轉文字
def continuous_recognition(path):
    audio_config = speechsdk.audio.AudioConfig(filename=path)
    speech_config = speechsdk.SpeechConfig(subscription="azure金鑰", region="westus2")
    speech_config.speech_recognition_language="zh-tw"
    speech_recognizer = speechsdk.SpeechRecognizer(speech_config=speech_config, audio_config=audio_config)
    #phrase_list_grammar = speechsdk.PhraseListGrammar.from_recognizer(speech_recognizer)
    #phrase_list_grammar.addPhrase(["自定義字詞"])
    done = False
    def stop_cb(evt):
        print('CLOSING on {}'.format(evt))
        speech_recognizer.stop_continuous_recognition()
        nonlocal done
        done = True
    text=[]
    speech_recognizer.recognized.connect(lambda evt: text.append(evt.result.text))
    speech_recognizer.session_started.connect(lambda evt: print('SESSION STARTED: {}'.format(evt['text'])))
    speech_recognizer.session_stopped.connect(lambda evt: print('SESSION STOPPED {}'.format(evt)))
    
    speech_recognizer.session_stopped.connect(stop_cb)
    speech_recognizer.canceled.connect(stop_cb)
    
    speech_recognizer.start_continuous_recognition()
    while not done:
        time.sleep(.5)
    print("第二步：语音转文本完成")
    sg.popup('語音轉文字完成')
    print(text)
    return text
text=continuous_recognition(file_name+".wav")


#開始整理成Word檔案
sg.popup('開始整理word檔案......')

#在word裡面建立逐字稿和經過關鍵字標記的逐字稿
document = Document()
document.add_heading('會議記錄逐字稿未經過標記', level=1)
document.add_paragraph(text)
document.save('docx_file.docx')
doc = Document('docx_file.docx')
p1_text = doc.paragraphs[1].text
doc.add_page_break()
doc.add_page_break()
doc.add_heading('會議記錄逐字稿標記後', level=1)
p2 = doc.add_paragraph()
substrings = p1_text.split(values[1][1])
for substring in substrings[:-1]:
    p2.add_run(substring)
    font = p2.add_run(values[1][1]).font
    font.highlight_color = WD_COLOR_INDEX.YELLOW
p2.add_run(substrings[-1])

#Word檔案儲存
doc.save(values[1][0]+'檔案名稱.docx')



